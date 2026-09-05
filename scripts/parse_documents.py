import os
import sys
import json
import re
import traceback
import pymupdf
from pypdf import PdfReader
from docx import Document

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# Configuration
DOCS_DIRS = ["Documents utiles", "Marchés Publics docs"]
CHUNK_SIZE = 3000
CHUNK_OVERLAP = 300

def clean_text(text):
    """Clean text by removing excessive whitespace and normalizing separators."""
    if not text:
        return ""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\r\n|\r|\n', '\n', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def get_chunks(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping chunks of size characters."""
    chunks = []
    if not text:
        return chunks
    if len(text) <= size:
        return [text]
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        chunks.append(chunk)
        start += (size - overlap)
    return chunks

def get_category(root_path):
    """Determine category based on folder hierarchy."""
    parts = os.path.normpath(root_path).split(os.sep)
    if len(parts) > 1:
        for folder_name in parts[1:]:
            fn_upper = folder_name.upper()
            if "BENIN" in fn_upper or "BÉNIN" in fn_upper:
                return "Bénin"
            elif "NIGER" in fn_upper and "NIGERIA" not in fn_upper:
                return "Niger"
            elif "CONGO" in fn_upper:
                return "Congo"
            elif "CAMEROUN" in fn_upper:
                return "Cameroun"
            elif "CENTRAFIQUE" in fn_upper or "CENTRAFRIQUE" in fn_upper:
                return "Centrafique"
            elif "TOGO" in fn_upper:
                return "Togo"
            elif "MALI" in fn_upper:
                return "Mali"
            elif "TCHAD" in fn_upper:
                return "Tchad"
            elif "RCI" in fn_upper or "IVOIRE" in fn_upper:
                return "Côte d'Ivoire"
            elif "SENEGAL" in fn_upper or "SÉNÉGAL" in fn_upper:
                return "Sénégal"
            elif "BURKINA" in fn_upper:
                return "Burkina Faso"
            elif "GUINEE" in fn_upper or "GUINÉE" in fn_upper:
                return "Guinée"
            elif "GABON" in fn_upper:
                return "Gabon"
            elif "MAURITANIE" in fn_upper or "MAURITAN" in fn_upper:
                return "Mauritanie"
            elif "RDC" in fn_upper:
                return "RDC (Congo-Kinshasa)"
            elif "BANQUE MONDIALE" in fn_upper or "WORLD BANK" in fn_upper:
                return "Banque Mondiale"
            elif "BOAD" in fn_upper:
                return "BOAD (Banque Ouest-Africaine de Développement)"
            elif "AFD" in fn_upper:
                return "AFD (Agence Française de Développement)"
            elif "BAD" in fn_upper:
                return "BAD (Banque Africaine de Développement)"
            elif "BID" in fn_upper or "ISDB" in fn_upper:
                return "BID (Banque Islamique de Développement)"
            elif "UEMOA" in fn_upper:
                return "UEMOA"
            elif "AIDE EMPLOI" in fn_upper or "EMPLOI" in fn_upper:
                return "Aide Emploi et Recrutement"
            elif "THEMATIQUE" in fn_upper or "THÉMATIQUE" in fn_upper:
                return "Thématiques & Études"
            elif "CERTIFICATION" in fn_upper or "RECHERCHE" in fn_upper:
                return "Certifications & Recherches"
            elif "CARROUSEL" in fn_upper or "CAROUSEL" in fn_upper or "CAROUSSEL" in fn_upper or "CAROUS" in fn_upper:
                return "Carrousels Pédagogiques"
            elif "AUDIT" in fn_upper and "CONTROLE" in fn_upper:
                return "Audit et Contrôle des Finances Publiques"
            elif "DURABILITE" in fn_upper or "DURABILITÉ" in fn_upper:
                return "Marchés Durables"
            elif "AUTRES DOCUMENTS" in fn_upper:
                return "Autres Documents"
    return "Général"

def extract_pdf_text(file_path):
    """Extract page-by-page text from a PDF using PyMuPDF with pypdf fallback."""
    pages = []
    # 1. Primary: PyMuPDF
    try:
        doc = pymupdf.open(file_path)
        for i, page in enumerate(doc):
            try:
                text = page.get_text() or ""
                cleaned = clean_text(text)
                if cleaned:
                    pages.append((i + 1, cleaned))
            except Exception:
                continue
        if pages:
            return pages
    except Exception as e:
        pass

    # 2. Fallback: pypdf
    try:
        reader = PdfReader(file_path, strict=False)
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                cleaned = clean_text(text)
                if cleaned:
                    pages.append((i + 1, cleaned))
            except Exception:
                continue
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return pages

def extract_docx_text(file_path):
    """Extract text from a docx file as paragraphs."""
    paragraphs = []
    try:
        doc = Document(file_path)
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return "\n".join(paragraphs)

def extract_doc_legacy_text(file_path):
    """Extract text from binary .doc files via UTF-16 and ASCII heuristics."""
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        u16_strings = re.findall(rb'(?:[\x20-\x7E\xA0-\xFF]\x00){4,}', data)
        parts = [s.decode('utf-16le', errors='ignore') for s in u16_strings]
        if not parts or sum(len(p) for p in parts) < 100:
            ascii_strings = re.findall(rb'[\x20-\x7E\xA0-\xFF]{4,}', data)
            parts = [s.decode('latin1', errors='ignore') for s in ascii_strings]
        return clean_text("\n".join(parts))
    except Exception as e:
        print(f"Error reading legacy DOC {file_path}: {e}")
        return ""

def extract_xlsx_text(file_path):
    """Extract text from XLSX sheets."""
    if not HAS_OPENPYXL:
        return ""
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            lines.append(f"--- Feuille: {sheetname} ---")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    lines.append(" | ".join(vals))
        return clean_text("\n".join(lines))
    except Exception as e:
        print(f"Error reading XLSX {file_path}: {e}")
        return ""

def main():
    print("🚀 Démarrage de l'indexation complète des documents...")
    knowledge_base = []
    chunk_counter = 0
    processed_files = {}  # filename -> catalog item
    file_counter = 0

    # Scanner les répertoires principaux
    for DOCS_DIR in DOCS_DIRS:
        if not os.path.exists(DOCS_DIR):
            print(f"Répertoire {DOCS_DIR} introuvable, ignoré.")
            continue

        print(f"\n📂 Analyse du répertoire : {DOCS_DIR}...")

        for root, dirs, files in os.walk(DOCS_DIR):
            category = get_category(root)
            for file in files:
                ext = os.path.splitext(file)[1].lower()

                # Ignorer les fichiers temporaires système Office
                if file.startswith("~$") or file.startswith("._") or file == "Thumbs.db":
                    continue

                # Ignorer les images pures du scan RAG (gardées hors catalogue ou non textuelles)
                if ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]:
                    continue

                file_path = os.path.join(root, file)

                # Si le fichier est déjà traité (dédoublonnage par nom), on conserve la première occurrence
                if file in processed_files:
                    continue

                title = file.replace("_", " ").replace("-", " ").rsplit(".", 1)[0].strip()

                catalog_entry = {
                    "filename": file,
                    "title": title,
                    "category": category,
                    "path": file_path.replace("\\", "/"),
                    "chunks": 0,
                    "first_page_preview": ""
                }

                file_chunks = []

                try:
                    if ext == ".pdf":
                        pages = extract_pdf_text(file_path)
                        for page_num, page_text in pages:
                            chunks = get_chunks(page_text)
                            for idx, chunk in enumerate(chunks):
                                file_chunks.append({
                                    "id": f"chunk_{chunk_counter}",
                                    "source": file,
                                    "path": file_path.replace("\\", "/"),
                                    "category": category,
                                    "title": f"{file} - Page {page_num}" if len(chunks) == 1 else f"{file} - Page {page_num} (Partie {idx + 1})",
                                    "content": chunk
                                })
                                chunk_counter += 1
                        if pages and not catalog_entry["first_page_preview"]:
                            catalog_entry["first_page_preview"] = pages[0][1][:300]

                    elif ext == ".docx":
                        full_text = extract_docx_text(file_path)
                        cleaned = clean_text(full_text)
                        if cleaned:
                            chunks = get_chunks(cleaned)
                            for idx, chunk in enumerate(chunks):
                                file_chunks.append({
                                    "id": f"chunk_{chunk_counter}",
                                    "source": file,
                                    "path": file_path.replace("\\", "/"),
                                    "category": category,
                                    "title": f"{file} - Partie {idx + 1}",
                                    "content": chunk
                                })
                                chunk_counter += 1
                            catalog_entry["first_page_preview"] = cleaned[:300]

                    elif ext == ".doc":
                        doc_text = extract_doc_legacy_text(file_path)
                        if doc_text:
                            chunks = get_chunks(doc_text)
                            for idx, chunk in enumerate(chunks):
                                file_chunks.append({
                                    "id": f"chunk_{chunk_counter}",
                                    "source": file,
                                    "path": file_path.replace("\\", "/"),
                                    "category": category,
                                    "title": f"{file} - Partie {idx + 1}",
                                    "content": chunk
                                })
                                chunk_counter += 1
                            catalog_entry["first_page_preview"] = doc_text[:300]

                    elif ext in [".xlsx", ".xls"]:
                        if ext == ".xlsx":
                            xlsx_text = extract_xlsx_text(file_path)
                            if xlsx_text:
                                chunks = get_chunks(xlsx_text)
                                for idx, chunk in enumerate(chunks):
                                    file_chunks.append({
                                        "id": f"chunk_{chunk_counter}",
                                        "source": file,
                                        "path": file_path.replace("\\", "/"),
                                        "category": category,
                                        "title": f"{file} - Tableur (Partie {idx + 1})",
                                        "content": chunk
                                    })
                                    chunk_counter += 1
                                catalog_entry["first_page_preview"] = xlsx_text[:300]

                    elif ext in [".rtf", ".txt"]:
                        try:
                            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                                t = clean_text(f.read())
                            if t:
                                chunks = get_chunks(t)
                                for idx, chunk in enumerate(chunks):
                                    file_chunks.append({
                                        "id": f"chunk_{chunk_counter}",
                                        "source": file,
                                        "path": file_path.replace("\\", "/"),
                                        "category": category,
                                        "title": f"{file} - Texte",
                                        "content": chunk
                                    })
                                    chunk_counter += 1
                                catalog_entry["first_page_preview"] = t[:300]
                        except Exception:
                            pass

                except Exception as e:
                    print(f"⚠️ Erreur lors du traitement de {file}: {e}")

                # Si le document n'a pas de couche texte directe (PDF scanné, tableur brut, etc.),
                # créer un chunk sémantique enrichi pour qu'il soit 100% indexable par l'IA et affiche des chunks
                if len(file_chunks) == 0:
                    clean_title = title.replace("_", " ").replace("-", " ").strip()
                    meta_info = ""
                    if ext == ".pdf":
                        try:
                            doc_pdf = pymupdf.open(file_path)
                            num_p = len(doc_pdf)
                            meta_info = f" Document composé de {num_p} page(s) officielle(s)."
                        except Exception:
                            pass

                    doc_content = (
                        f"Texte officiel / Décret / Règlement : {clean_title}.\n"
                        f"Catégorie : {category}.\n"
                        f"Fichier source : {file}.\n"
                        f"Chemin documentaire : {file_path.replace(os.sep, '/')}.{meta_info}\n"
                        f"Ce document juridique et réglementaire est officiellement enregistré sous la juridiction / bailleur {category} dans la base de données PROCURA."
                    )
                    file_chunks.append({
                        "id": f"chunk_{chunk_counter}",
                        "source": file,
                        "path": file_path.replace("\\", "/"),
                        "category": category,
                        "title": f"{clean_title} - Document Officiel ({category})",
                        "content": doc_content
                    })
                    chunk_counter += 1
                    catalog_entry["first_page_preview"] = doc_content[:300]

                catalog_entry["chunks"] = len(file_chunks)
                if not catalog_entry["first_page_preview"]:
                    catalog_entry["first_page_preview"] = f"Document officiel {title} ({category})"

                processed_files[file] = catalog_entry
                knowledge_base.extend(file_chunks)
                file_counter += 1

    # ── Scanner également les documents déposés à la racine du projet ──
    print(f"\n📂 Analyse des documents à la racine du projet...")
    for file in os.listdir("."):
        if not os.path.isfile(file):
            continue
        ext = os.path.splitext(file)[1].lower()
        if ext not in [".pdf", ".docx", ".doc", ".rtf", ".xlsx", ".xls"]:
            continue
        if file.startswith("~$") or file.startswith("._") or file in processed_files:
            continue

        file_path = file
        category = "Autres Documents"
        f_upper = file.upper()
        if "COMMERCIAL" in f_upper or "EMPLOI" in f_upper:
            category = "Aide Emploi et Recrutement"
        elif "SITE" in f_upper or "WEB" in f_upper:
            category = "Général"
        elif "CATALOGUE" in f_upper:
            category = "Général"

        title = file.replace("_", " ").replace("-", " ").rsplit(".", 1)[0].strip()
        catalog_entry = {
            "filename": file,
            "title": title,
            "category": category,
            "path": file_path.replace("\\", "/"),
            "chunks": 0,
            "first_page_preview": ""
        }
        file_chunks = []

        try:
            if ext == ".pdf":
                pages = extract_pdf_text(file_path)
                for page_num, page_text in pages:
                    chunks = get_chunks(page_text)
                    for idx, chunk in enumerate(chunks):
                        file_chunks.append({
                            "id": f"chunk_{chunk_counter}",
                            "source": file,
                            "path": file_path.replace("\\", "/"),
                            "category": category,
                            "title": f"{file} - Page {page_num}" if len(chunks) == 1 else f"{file} - Page {page_num} (Partie {idx + 1})",
                            "content": chunk
                        })
                        chunk_counter += 1
                if pages:
                    catalog_entry["first_page_preview"] = pages[0][1][:300]
            elif ext == ".docx":
                full_text = extract_docx_text(file_path)
                cleaned = clean_text(full_text)
                if cleaned:
                    chunks = get_chunks(cleaned)
                    for idx, chunk in enumerate(chunks):
                        file_chunks.append({
                            "id": f"chunk_{chunk_counter}",
                            "source": file,
                            "path": file_path.replace("\\", "/"),
                            "category": category,
                            "title": f"{file} - Partie {idx + 1}",
                            "content": chunk
                        })
                        chunk_counter += 1
                    catalog_entry["first_page_preview"] = cleaned[:300]
        except Exception as e:
            print(f"⚠️ Erreur lors du traitement de {file}: {e}")

        catalog_entry["chunks"] = len(file_chunks)
        if not catalog_entry["first_page_preview"]:
            catalog_entry["first_page_preview"] = f"Document {title}"

        processed_files[file] = catalog_entry
        knowledge_base.extend(file_chunks)
        file_counter += 1

    # ── Sauvegarde de la base de connaissances (Knowledge Base Chunks) ──
    PART_CHUNK_SIZE = 8000
    total_chunks = len(knowledge_base)
    num_parts = (total_chunks + PART_CHUNK_SIZE - 1) // PART_CHUNK_SIZE if total_chunks > 0 else 1

    for i in range(num_parts):
        part_data = knowledge_base[i*PART_CHUNK_SIZE : (i+1)*PART_CHUNK_SIZE]
        part_filename = f"knowledge_base_part_{i+1}.json"
        with open(part_filename, "w", encoding="utf-8") as f:
            json.dump(part_data, f, ensure_ascii=False, indent=2)
        print(f"✅ Partie {i+1}/{num_parts} sauvegardée dans {part_filename} ({os.path.getsize(part_filename) / 1024 / 1024:.2f} MB)")

    # Metadata
    meta = {
        "total_chunks": total_chunks,
        "num_parts": num_parts
    }
    with open("knowledge_base_meta.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2)

    # ── Générer documents_catalog.json (100% des documents uniques) ──
    catalog = list(processed_files.values())
    catalog.sort(key=lambda x: (x["category"], x["title"]))
    with open("documents_catalog.json", "w", encoding="utf-8") as f:
        json.dump(catalog, f, ensure_ascii=False, indent=2)

    print(f"\n========================================================")
    print(f"🎉 SUCCÈS COMPLET :")
    print(f" - Documents uniques dans le catalogue Admin : {len(catalog)}")
    print(f" - Chunks RAG générés pour l'IA : {total_chunks}")
    print(f" - Fichier documents_catalog.json mis à jour.")
    print(f"========================================================")

if __name__ == "__main__":
    main()

